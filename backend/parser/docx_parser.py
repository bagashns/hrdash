"""
DOCX text extraction (python-docx).

Membaca paragraf + isi tabel, karena banyak CV menempatkan bagian
"Pendidikan" / "Pengalaman" di dalam tabel dua kolom.
"""
from pathlib import Path

from docx import Document


def extract_text_from_docx(path: Path) -> str:
    doc = Document(path)

    chunks: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            chunks.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                chunks.append(row_text)

    return "\n".join(chunks).strip()
